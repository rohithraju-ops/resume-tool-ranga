"""Builds a Word doc from tailored_data, mirroring the PDF layout."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_horizontal_rule(paragraph):
    """Adds a bottom border to a paragraph (section underline)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _strip_latex(text):
    """Remove LaTeX artifacts from bullet text for clean docx rendering."""
    if not text:
        return ""
    # Drop \item markers
    text = re.sub(r'\\item\s*', '', text)
    # Drop \textbf{...} but keep inner content (we'll handle bold separately)
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    # Drop noindent, hfill, begin/end, topsep
    text = re.sub(r'\\noindent\s*', '', text)
    text = re.sub(r'\\hfill\s*', '\t', text)
    text = re.sub(r'\\begin\{itemize\}[^\n]*\n?', '', text)
    text = re.sub(r'\\end\{itemize\}\s*', '', text)
    # Unescape LaTeX specials back to readable form
    text = text.replace('\\&', '&').replace('\\%', '%').replace('\\$', '$').replace('\\#', '#')
    text = text.replace('\\_', '_')
    return text.strip()


def _split_bullets(bullets_block):
    """Split a block of \\item lines into individual bullet strings."""
    if not bullets_block:
        return []
    items = re.split(r'\\item\s+', bullets_block)
    return [_strip_latex(i).strip() for i in items if i.strip()]


def _parse_project_block(block):
    """
    Parse a project_N_block string into (name, institution, [bullets]).
    Expected format:
    \\noindent \\textbf{Project Name}, \\textit{Institution}
    \\begin{itemize}[topsep=1pt]
    \\item Bullet one.
    \\item Bullet two.
    ...
    \\end{itemize}
    """
    # Extract name from \textbf{...}
    name_match = re.search(r'\\textbf\{([^}]+)\}', block)
    name = name_match.group(1) if name_match else "Project"

    # Extract institution from \textit{...}
    inst_match = re.search(r'\\textit\{([^}]+)\}', block)
    institution = inst_match.group(1) if inst_match else ""

    # Extract bullets after \begin{itemize}
    body_match = re.search(r'\\begin\{itemize\}[^\n]*\n(.*?)\\end\{itemize\}', block, re.DOTALL)
    bullets = []
    if body_match:
        bullets = _split_bullets(body_match.group(1))

    return name, institution, bullets


def _set_font(run, name="Times New Roman", size=10, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # Force font for east-asian + ascii runs (Windows Word compatibility)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    rfonts.set(qn('w:cs'), name)


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, size=11, bold=True)
    _add_horizontal_rule(p)
    return p


def _add_right_aligned_line(doc, left_text, left_bold, right_text, right_italic=True, size=10):
    """Paragraph with left-aligned text and right-aligned text (via tab stop)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Right tab stop at ~7.3" (A4 text width)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)

    run_left = p.add_run(left_text)
    _set_font(run_left, size=size, bold=left_bold)

    if right_text:
        p.add_run("\t")
        run_right = p.add_run(right_text)
        _set_font(run_right, size=size, bold=False, italic=right_italic)
    return p


def _add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)

    # If text contains a tab (from \hfill), split for right-align
    if '\t' in text:
        left, right = text.split('\t', 1)
        run = p.add_run(left.strip())
        _set_font(run, size=size)
    else:
        run = p.add_run(text)
        _set_font(run, size=size)
    return p


def build_resume_docx(tailored_data: dict, output_path: str, profile: dict):
    """
    Build a .docx resume from the same tailored_data used to build the PDF.
    Matches the visual structure of resume_template.tex as closely as Word allows.
    """
    doc = Document()

    # Page margins to match template (~0.1in but Word minimum ~0.25in is safer)
    section = doc.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    # --- HEADER ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(profile.get("full_name", "").upper())
    _set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    contact = f"{profile.get('email', '')} | {profile.get('phone', '')} | {profile.get('linkedin', '')}"
    run = p.add_run(contact)
    _set_font(run, size=10)

    # --- SUMMARY ---
    _add_section_header(doc, "Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(tailored_data.get("summary", ""))
    _set_font(run, size=10)

    # --- EDUCATION AND CERTIFICATIONS ---
    _add_section_header(doc, "Education")
    _add_right_aligned_line(doc, "Master of Science in Mechanical Engineering", True, "May 2026")
    _add_right_aligned_line(doc, "Arizona State University, Arizona, USA", False, "GPA: 3.70")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run_bold = p.add_run("Relevant Coursework: ")
    _set_font(run_bold, size=10, bold=True)
    run = p.add_run("Modern Manufacturing Methods, Intermediate Thermodynamics")
    _set_font(run, size=10)

    _add_right_aligned_line(doc, "Bachelor of Technology in Mechanical Engineering", True, "May 2021")
    _add_right_aligned_line(doc, "Dayananda Sagar University, Karnataka, India", False, "CGPA: 8.36")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run_bold = p.add_run("Additional Coursework: ")
    _set_font(run_bold, size=10, bold=True)
    run = p.add_run("GD&T based on ASME Y14.5 (Udemy), Lean Six Sigma Yellow Belt (Udemy)")
    _set_font(run, size=10)

    # --- TECHNICAL SKILLS ---
    _add_section_header(doc, "Technical Skills")
    skills_raw = tailored_data.get("skills_block", "")
    # Each \item is a category line: \item \textbf{Category:} item1, item2, ...
    for item in re.split(r'\\item\s+', skills_raw):
        if not item.strip():
            continue
        # Match \textbf{Label:} content
        m = re.match(r'\\textbf\{([^}]+)\}\s*(.*)', item.strip(), re.DOTALL)
        if m:
            label, content = m.group(1), m.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            run_label = p.add_run(label + " ")
            _set_font(run_label, size=10, bold=True)
            run_rest = p.add_run(_strip_latex(content))
            _set_font(run_rest, size=10)
        else:
            _add_bullet(doc, _strip_latex(item))

    # --- PROFESSIONAL EXPERIENCE ---
    _add_section_header(doc, "Professional Experience")

    experiences = [
        ("Design Engineer, Aktis Engineering Solutions Pvt. Ltd., Bangalore, India", "Sep 2022 to Jul 2024", "exp_1_bullets"),
        ("Project Associate, Pactera EDGE (now Centific), Hyderabad, India", "Mar 2022 to Aug 2022", "exp_2_bullets"),
        ("Intern, Bosch Rexroth in Collaboration with DSU, Bangalore, Karnataka, India", "Feb 2021 to Apr 2021", "exp_3_bullets"),
    ]

    for idx, (heading, dates, field) in enumerate(experiences):
        _add_right_aligned_line(doc, heading, True, dates)
        for bullet in _split_bullets(tailored_data.get(field, "")):
            _add_bullet(doc, bullet)
        if idx < len(experiences) - 1:
            # small gap between jobs
            gap = doc.add_paragraph()
            gap.paragraph_format.space_before = Pt(0)
            gap.paragraph_format.space_after = Pt(2)

    # --- ACADEMIC PROJECTS ---
    _add_section_header(doc, "Academic Projects")
    for i in range(1, 6):
        block = tailored_data.get(f"project_{i}_block", "")
        if not block:
            continue
        name, institution, bullets = _parse_project_block(block)
        # Project heading: name in bold, comma, institution in italics
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(name)
        _set_font(run, size=10, bold=True)
        if institution:
            run_sep = p.add_run(", ")
            _set_font(run_sep, size=10)
            run_inst = p.add_run(institution)
            _set_font(run_inst, size=10, italic=True)
        for bullet in bullets:
            _add_bullet(doc, bullet)

    doc.save(output_path)